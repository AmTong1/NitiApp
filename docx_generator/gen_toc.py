# -*- coding: utf-8 -*-
import sys, io, glob
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn

files = glob.glob(r"C:\Users\tanut\Downloads\*V1.9.docx")
files = [f for f in files if '~$' not in f]
filepath = files[0]
print(f"Source: {filepath}")

doc = Document(filepath)

# === Collect figures & tables ===
figures = []
tables = []
for para in doc.paragraphs:
    t = para.text.strip()
    if t.startswith("รูปที่") or t.startswith("รูปที "):
        figures.append(t)
    if t.startswith("ตารางที่") or t.startswith("ตารางที "):
        tables.append(t)

print(f"Figures: {len(figures)}, Tables: {len(tables)}")

# === Collect headings ===
headings = []
for para in doc.paragraphs:
    t = para.text.strip()
    s = para.style.name if para.style else ""
    if not t:
        continue
    if s == "Heading 1":
        headings.append((1, t))
    elif s == "Heading 2":
        # Skip table captions that are incorrectly styled as Heading 2
        if t.startswith("ตารางที"):
            continue
        headings.append((2, t))
    elif s == "Heading 3":
        headings.append((3, t))
    elif s == "Normal" and any(kw in t for kw in ['บทคัดย่อ', 'กิตติกรรม', 'สารบัญ', 'บรรณานุกรม']):
        if t in ['บทคัดย่อ', 'กิตติกรรมประกาศ', 'สารบัญ', 'บรรณานุกรม']:
            headings.append((0, t))  # level 0 = front matter

# Build new document
newdoc = Document()

# === Page setup ===
for section in newdoc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.81)
    section.right_margin = Cm(2.54)

def set_font(run, size=16, bold=False):
    run.font.name = 'TH Sarabun PSK'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:cs'), 'TH Sarabun PSK')

def add_toc_entry(doc, text, level=0, page=""):
    """Add a TOC entry with dot leader tab to right margin"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    
    # Set left indent based on level
    if level == 0:
        p.paragraph_format.left_indent = Cm(0)
    elif level == 1:
        p.paragraph_format.left_indent = Cm(0)
    elif level == 2:
        p.paragraph_format.left_indent = Cm(1.0)
    elif level == 3:
        p.paragraph_format.left_indent = Cm(2.0)
    
    # Add tab stop with dot leader at right margin
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(13.65), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    
    run = p.add_run(text)
    bold = level <= 1
    set_font(run, 16, bold)
    
    # Add tab + page number
    tab_run = p.add_run("\t")
    set_font(tab_run, 16, False)
    
    page_run = p.add_run(page)
    set_font(page_run, 16, False)

def add_centered_title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_font(run, size, True)

def add_blank_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("")
    set_font(run, 16, False)

# =====================================================
# 1. สารบัญ (Table of Contents)
# =====================================================
add_centered_title(newdoc, "สารบัญ")

# Add "เรื่อง" and "หน้า" headers
p_header = newdoc.add_paragraph()
p_header.paragraph_format.space_before = Pt(0)
p_header.paragraph_format.space_after = Pt(6)
tab_stops = p_header.paragraph_format.tab_stops
tab_stops.add_tab_stop(Cm(13.65), WD_TAB_ALIGNMENT.RIGHT)
run1 = p_header.add_run("เรื่อง")
set_font(run1, 16, True)
run_tab = p_header.add_run("\t")
set_font(run_tab, 16, False)
run2 = p_header.add_run("หน้า")
set_font(run2, 16, True)

# Front matter
front_items = [
    ("บทคัดย่อ", ""),
    ("ABSTRACT", ""),
    ("กิตติกรรมประกาศ", ""),
    ("สารบัญ", ""),
    ("สารบัญตาราง", ""),
    ("สารบัญรูป", ""),
]
for text, page in front_items:
    add_toc_entry(newdoc, text, 0, page)

add_blank_line(newdoc)

# Chapter 1
add_toc_entry(newdoc, "บทที่ 1 บทนำ", 1, "")
ch1_sections = [
    (2, "1.1 ที่มาและความสำคัญของปัญหา"),
    (2, "1.2 วัตถุประสงค์ของโครงงาน"),
    (2, "1.3 ขอบเขตของโครงงาน"),
    (2, "1.4 แผนดำเนินการและระยะเวลาในการดำเนินโครงงาน"),
    (2, "1.5 ประโยชน์ที่คาดว่าจะได้รับ"),
    (2, "1.6 งบประมาณที่ใช้ในการจัดทำโครงงาน"),
]
for lv, text in ch1_sections:
    add_toc_entry(newdoc, text, lv, "")

add_blank_line(newdoc)

# Chapter 2
add_toc_entry(newdoc, "บทที่ 2 ทฤษฎีและงานวิจัยที่เกี่ยวข้อง", 1, "")
ch2_sections = [
    (2, "2.1 ทฤษฎีด้านการพัฒนาส่วนติดต่อผู้ใช้งาน (Frontend Technology)"),
    (2, "2.2 ทฤษฎีด้านการประมวลผลและการสื่อสาร (Backend & Communication)"),
    (2, "2.3 ระบบจัดการฐานข้อมูลและความปลอดภัย (Database & Security)"),
    (2, "2.4 ระบบบริการภายนอกและการเงินดิจิทัล"),
    (2, "2.5 เครื่องมือสำหรับการพัฒนาแอปพลิเคชัน"),
    (2, "2.6 งานวิจัยที่เกี่ยวข้อง"),
]
for lv, text in ch2_sections:
    add_toc_entry(newdoc, text, lv, "")

add_blank_line(newdoc)

# Chapter 3
add_toc_entry(newdoc, "บทที่ 3 วิธีการดำเนินงาน", 1, "")
ch3_sections = [
    (2, "3.1 เครื่องมือที่ใช้ในการพัฒนาระบบ"),
    (2, "3.2 ขั้นตอนการดำเนินงาน"),
    (2, "3.3 การออกแบบการทำงานของระบบ"),
    (2, "3.4 การกำหนดตัวแปรและโครงสร้างข้อมูล (Data Dictionary)"),
    (2, "3.5 การออกแบบส่วนติดต่อประสานงานกับผู้ใช้"),
    (2, "3.6 การออกแบบฐานข้อมูล"),
]
for lv, text in ch3_sections:
    add_toc_entry(newdoc, text, lv, "")

add_blank_line(newdoc)

# Chapter 4
add_toc_entry(newdoc, "บทที่ 4 ผลการดำเนินงาน", 1, "")
ch4_sections = [
    (2, "4.1 ผลของการพัฒนาระบบ"),
    (2, "4.2 ผลของการทดสอบความถูกต้องของลิงก์และการนำทาง (Link & Navigation Testing)"),
    (2, "4.3 ผลการทดสอบการยอมรับของผู้ใช้งาน (UAT)"),
    (2, "4.4 ผลการประเมินประสิทธิภาพและความพึงพอใจ"),
]
for lv, text in ch4_sections:
    add_toc_entry(newdoc, text, lv, "")

add_blank_line(newdoc)

# Chapter 5
add_toc_entry(newdoc, "บทที่ 5 สรุปผล อภิปรายผล และข้อเสนอแนะ", 1, "")
ch5_sections = [
    (2, "5.1 สรุปผลการดำเนินงาน (Conclusion)"),
    (2, "5.2 อภิปรายผล (Discussion)"),
    (2, "5.3 ปัญหาและอุปสรรค (Problems and Obstacles)"),
    (2, "5.4 ข้อเสนอแนะเพื่อการพัฒนาต่อยอด (Recommendations for Future Work)"),
]
for lv, text in ch5_sections:
    add_toc_entry(newdoc, text, lv, "")

add_blank_line(newdoc)
add_toc_entry(newdoc, "บรรณานุกรม", 0, "")

# =====================================================
# Page break before สารบัญตาราง
# =====================================================
newdoc.add_page_break()
add_centered_title(newdoc, "สารบัญตาราง")

p_header2 = newdoc.add_paragraph()
p_header2.paragraph_format.space_before = Pt(0)
p_header2.paragraph_format.space_after = Pt(6)
tab_stops2 = p_header2.paragraph_format.tab_stops
tab_stops2.add_tab_stop(Cm(13.65), WD_TAB_ALIGNMENT.RIGHT)
run1 = p_header2.add_run("ตาราง")
set_font(run1, 16, True)
run_tab = p_header2.add_run("\t")
set_font(run_tab, 16, False)
run2 = p_header2.add_run("หน้า")
set_font(run2, 16, True)

for t_text in tables:
    add_toc_entry(newdoc, t_text, 2, "")

# =====================================================
# Page break before สารบัญรูป
# =====================================================
newdoc.add_page_break()
add_centered_title(newdoc, "สารบัญรูป")

p_header3 = newdoc.add_paragraph()
p_header3.paragraph_format.space_before = Pt(0)
p_header3.paragraph_format.space_after = Pt(6)
tab_stops3 = p_header3.paragraph_format.tab_stops
tab_stops3.add_tab_stop(Cm(13.65), WD_TAB_ALIGNMENT.RIGHT)
run1 = p_header3.add_run("รูป")
set_font(run1, 16, True)
run_tab = p_header3.add_run("\t")
set_font(run_tab, 16, False)
run2 = p_header3.add_run("หน้า")
set_font(run2, 16, True)

for f_text in figures:
    add_toc_entry(newdoc, f_text, 2, "")

# Save
output_path = r"C:\Users\tanut\Downloads\สารบัญ_NitiSmart_V1.9.docx"
newdoc.save(output_path)
print(f"\nSaved to: {output_path}")
print("Done!")
