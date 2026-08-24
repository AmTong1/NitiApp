# -*- coding: utf-8 -*-
import sys, io, glob
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

files = glob.glob(r"C:\Users\tanut\Downloads\*V1.9.docx")
files = [f for f in files if '~$' not in f]
filepath = files[0]
doc = Document(filepath)

print("=== FIGURES (รูปที่) ===")
for para in doc.paragraphs:
    t = para.text.strip()
    if t.startswith("รูปที่") or t.startswith("รูปที"):
        print(f"  {t}")

print("\n=== TABLES (ตารางที่) ===")
for para in doc.paragraphs:
    t = para.text.strip()
    if t.startswith("ตารางที่") or t.startswith("ตารางที"):
        print(f"  {t}")

print("\n=== ALL HEADINGS ===")
for para in doc.paragraphs:
    t = para.text.strip()
    s = para.style.name if para.style else ""
    if 'Heading' in s and t:
        print(f"  [{s}] {t}")

print("\n=== KEY PARAGRAPHS ===")
for para in doc.paragraphs:
    t = para.text.strip()
    s = para.style.name if para.style else ""
    if t and any(kw in t for kw in ['บทที่', 'สารบัญ', 'บทคัดย่อ', 'ABSTRACT', 'กิตติกรรม', 'ภาคผนวก', 'บรรณานุกรม']):
        print(f"  [{s}] {t[:120]}")
