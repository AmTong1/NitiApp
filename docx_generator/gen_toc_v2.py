# -*- coding: utf-8 -*-
import sys, io, glob, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
import pdfplumber

# ========== 1) Read page numbers from PDF ==========
files = glob.glob(r"C:\Users\tanut\Downloads\*V1.7.pdf")
files = [f for f in files if '~$' not in f]
pdf = pdfplumber.open(files[0])

pg = {}  # key -> page number
for i, page in enumerate(pdf.pages):
    text = page.extract_text()
    if not text: continue
    pn = i + 1
    if 8 <= pn <= 9: continue  # skip TOC pages
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if line in ['บทคัดย่อ','ABSTRACT','กิตติกรรมประกาศ','บรรณานุกรม']:
            pg.setdefault(line, pn)
        if re.match(r'^บทที่\s*\d+$', line):
            pg.setdefault(line, pn)
        m = re.match(r'^(\d+\.\d+)\s', line)
        if m: pg.setdefault(m.group(1), pn)
        m = re.match(r'(รูปที่?\s*\d+\.\d+)', line)
        if m: pg.setdefault(m.group(1).strip(), pn)
        m = re.match(r'(ตารางที่?\s*[\d.]+)', line)
        if m: pg.setdefault(m.group(1).strip().rstrip(':'), pn)
pdf.close()
print(f"Mapped {len(pg)} items to pages")

def P(key):
    """Lookup page number"""
    for k,v in pg.items():
        if key in k or k in key: return str(v)
    return ""

# ========== 2) Read figure/table captions from docx ==========
dfiles = glob.glob(r"C:\Users\tanut\Downloads\*V1.9.docx")
dfiles = [f for f in dfiles if '~$' not in f]
doc = Document(dfiles[0])
figures, tables = [], []
for para in doc.paragraphs:
    t = para.text.strip()
    if t.startswith("รูปที่") or t.startswith("รูปที "): figures.append(t)
    if t.startswith("ตารางที่") or t.startswith("ตารางที "): tables.append(t)
print(f"Figures: {len(figures)}, Tables: {len(tables)}")

# ========== 3) Build TOC document ==========
newdoc = Document()
for s in newdoc.sections:
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.81); s.right_margin = Cm(2.54)

def sf(run, size=16, bold=False):
    run.font.name = 'TH Sarabun PSK'; run.font.size = Pt(size); run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:cs'), 'TH Sarabun PSK')

def toc(doc, text, level=0, page=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    indent = {0:0, 1:0, 2:1.0, 3:2.0}
    p.paragraph_format.left_indent = Cm(indent.get(level,0))
    p.paragraph_format.tab_stops.add_tab_stop(Cm(13.65), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(text); sf(r, 16, level<=1)
    t = p.add_run("\t"); sf(t, 16)
    r2 = p.add_run(page); sf(r2, 16)

def title(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text); sf(r, 18, True)

def header_row(doc, left, right):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(13.65), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(left); sf(r1, 16, True)
    p.add_run("\t")
    r2 = p.add_run(right); sf(r2, 16, True)

def blank(doc):
    p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
    sf(p.add_run(""), 16)

# --- สารบัญ ---
title(newdoc, "สารบัญ")
header_row(newdoc, "เรื่อง", "หน้า")

for t,p in [("บทคัดย่อ",P("บทคัดย่อ")),("ABSTRACT",P("ABSTRACT")),("กิตติกรรมประกาศ",P("กิตติกรรมประกาศ")),
            ("สารบัญ",""),("สารบัญตาราง",""),("สารบัญรูป","")]:
    toc(newdoc, t, 0, p)
blank(newdoc)

toc(newdoc, "บทที่ 1 บทนำ", 1, P("บทที่ 1"))
for s in ["1.1","1.2","1.3","1.4","1.5","1.6"]:
    names = {"1.1":"ที่มาและความสำคัญของปัญหา","1.2":"วัตถุประสงค์ของโครงงาน","1.3":"ขอบเขตของโครงงาน",
             "1.4":"แผนดำเนินการและระยะเวลา","1.5":"ประโยชน์ที่คาดว่าจะได้รับ","1.6":"งบประมาณที่ใช้ในการจัดทำโครงงาน"}
    toc(newdoc, f"{s} {names[s]}", 2, P(s))
blank(newdoc)

toc(newdoc, "บทที่ 2 ทฤษฎีและงานวิจัยที่เกี่ยวข้อง", 1, P("บทที่ 2"))
for s,n in [("2.1","ทฤษฎีด้านการพัฒนาส่วนติดต่อผู้ใช้งาน"),("2.2","ทฤษฎีด้านการประมวลผลและการสื่อสาร"),
            ("2.3","ระบบจัดการฐานข้อมูลและความปลอดภัย"),("2.4","ระบบบริการภายนอกและการเงินดิจิทัล"),
            ("2.5","เครื่องมือสำหรับการพัฒนาแอปพลิเคชัน"),("2.6","งานวิจัยที่เกี่ยวข้อง")]:
    toc(newdoc, f"{s} {n}", 2, P(s))
blank(newdoc)

toc(newdoc, "บทที่ 3 วิธีการดำเนินงาน", 1, P("บทที่ 3"))
for s,n in [("3.1","เครื่องมือที่ใช้ในการพัฒนาระบบ"),("3.2","ขั้นตอนการดำเนินงาน"),
            ("3.3","การออกแบบการทำงานของระบบ"),("3.4","การกำหนดตัวแปรและโครงสร้างข้อมูล"),
            ("3.5","การออกแบบส่วนติดต่อประสานงานกับผู้ใช้"),("3.6","การออกแบบฐานข้อมูล")]:
    toc(newdoc, f"{s} {n}", 2, P(s))
blank(newdoc)

toc(newdoc, "บทที่ 4 ผลการดำเนินงาน", 1, P("บทที่ 4"))
for s,n in [("4.1","ผลของการพัฒนาระบบ"),("4.2","ผลของการทดสอบความถูกต้องของลิงก์และการนำทาง"),
            ("4.3","ผลการทดสอบการยอมรับของผู้ใช้งาน (UAT)"),("4.4","ผลการประเมินประสิทธิภาพและความพึงพอใจ")]:
    toc(newdoc, f"{s} {n}", 2, P(s))
blank(newdoc)

toc(newdoc, "บทที่ 5 สรุปผล อภิปรายผล และข้อเสนอแนะ", 1, P("บทที่ 5"))
for s,n in [("5.1","สรุปผลการดำเนินงาน"),("5.2","อภิปรายผล"),
            ("5.3","ปัญหาและอุปสรรค"),("5.4","ข้อเสนอแนะเพื่อการพัฒนาต่อยอด")]:
    toc(newdoc, f"{s} {n}", 2, P(s))
blank(newdoc)
toc(newdoc, "บรรณานุกรม", 0, P("บรรณานุกรม"))

# --- สารบัญตาราง ---
newdoc.add_page_break()
title(newdoc, "สารบัญตาราง")
header_row(newdoc, "ตาราง", "หน้า")
for t in tables:
    m = re.match(r'(ตารางที่?\s*[\d.]+)', t)
    p = P(m.group(1).strip()) if m else ""
    toc(newdoc, t, 2, p)

# --- สารบัญรูป ---
newdoc.add_page_break()
title(newdoc, "สารบัญรูป")
header_row(newdoc, "รูป", "หน้า")
for f in figures:
    m = re.match(r'(รูปที่?\s*\d+\.\d+)', f)
    p = P(m.group(1).strip()) if m else ""
    toc(newdoc, f, 2, p)

out = r"C:\Users\tanut\Downloads\สารบัญ_NitiSmart_V2.docx"
newdoc.save(out)
print(f"Saved: {out}")
