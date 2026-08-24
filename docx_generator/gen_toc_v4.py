# -*- coding: utf-8 -*-
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn

# ========== 1) Load mapping from XML data ==========
raw_pg = {}
with io.open(r"d:\upgit\proj\docx_generator\xml_pages2.txt", "r", encoding='utf-8') as f:
    for line in f:
        line = line.strip()
        if not line: continue
        parts = line.split("::::")
        if len(parts) == 2:
            key, pn = parts[0], int(parts[1])
            if key not in raw_pg:
                raw_pg[key] = pn

OFFSET = 19

def P(key):
    # Lookup
    for k, v in raw_pg.items():
        if key == k or k.startswith(key):
            if v > OFFSET:
                return str(v - OFFSET)
    return ""

def P_exact(key):
    if key in raw_pg:
        v = raw_pg[key]
        if v > OFFSET: return str(v - OFFSET)
    # fallback
    return P(key)

# Get all figures and tables in order
figures = []
tables = []
with io.open(r"d:\upgit\proj\docx_generator\xml_pages2.txt", "r", encoding='utf-8') as f:
    for line in f:
        line = line.strip()
        if not line: continue
        parts = line.split("::::")
        if len(parts) == 2:
            k = parts[0]
            if k.startswith("รูปที่") or k.startswith("รูปที "):
                figures.append(k)
            if k.startswith("ตารางที่") or k.startswith("ตารางที "):
                tables.append(k)

# Deduplicate keeping order
def dedupe(seq):
    seen = set()
    return [x for x in seq if not (x in seen or seen.add(x))]

figures = dedupe(figures)
tables = dedupe(tables)

# ========== 2) Build TOC document ==========
newdoc = Document()
for s in newdoc.sections:
    s.top_margin=Cm(2.54); s.bottom_margin=Cm(2.54)
    s.left_margin=Cm(3.81); s.right_margin=Cm(2.54)

def sf(run, size=16, bold=False):
    run.font.name='TH Sarabun PSK'; run.font.size=Pt(size); run.font.bold=bold
    run._element.rPr.rFonts.set(qn('w:cs'),'TH Sarabun PSK')

def entry(doc, text, level=0, page=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(2)
    indent = {0:0,1:0,2:1.0,3:2.0}
    p.paragraph_format.left_indent = Cm(indent.get(level,0))
    p.paragraph_format.tab_stops.add_tab_stop(Cm(13.65), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r=p.add_run(text); sf(r,16, level<=1)
    p.add_run("\t"); r2=p.add_run(page); sf(r2,16)

def title(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after=Pt(12)
    r=p.add_run(text); sf(r,18,True)

def hdr(doc, left, right):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(13.65), WD_TAB_ALIGNMENT.RIGHT)
    r1=p.add_run(left); sf(r1,16,True)
    p.add_run("\t"); r2=p.add_run(right); sf(r2,16,True)

def blank(doc):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
    sf(p.add_run(""),16)

# --- สารบัญ ---
title(newdoc, "สารบัญ")
hdr(newdoc, "เรื่อง", "หน้า")

front = [("บทคัดย่อภาษาไทย","ก"),("บทคัดย่อภาษาอังกฤษ","ข"),("กิตติกรรมประกาศ","ค"),
         ("สารบัญ","ง"),("สารบัญตาราง","จ"),("สารบัญภาพ","ฉ")]
for t,p in front:
    entry(newdoc, t, 0, p)
blank(newdoc)

entry(newdoc, "บทที่ 1 บทนำ", 1, P_exact("บทที่ 1"))
for s,n in [("1.1","ที่มาและความสำคัญของปัญหา"),("1.2","วัตถุประสงค์ของโครงงาน"),
            ("1.3","ขอบเขตของโครงงาน"),("1.4","แผนดำเนินการและระยะเวลาในการดำเนินโครงงาน"),
            ("1.5","ประโยชน์ที่คาดว่าจะได้รับ"),("1.6","งบประมาณที่ใช้ในการจัดทำโครงงาน")]:
    entry(newdoc, f"{s} {n}", 2, P_exact(s))
blank(newdoc)

entry(newdoc, "บทที่ 2 ทฤษฎีและงานวิจัยที่เกี่ยวข้อง", 1, P_exact("บทที่ 2"))
for s,n in [("2.1","ทฤษฎีด้านการพัฒนาส่วนติดต่อผู้ใช้งาน"),("2.2","ทฤษฎีด้านการประมวลผลและการสื่อสาร"),
            ("2.3","ระบบจัดการฐานข้อมูลและความปลอดภัย"),("2.4","ระบบบริการภายนอกและการเงินดิจิทัล"),
            ("2.5","เครื่องมือสำหรับการพัฒนาแอปพลิเคชัน"),("2.6","งานวิจัยที่เกี่ยวข้อง")]:
    entry(newdoc, f"{s} {n}", 2, P_exact(s))
blank(newdoc)

entry(newdoc, "บทที่ 3 วิธีการดำเนินงาน", 1, P_exact("บทที่ 3"))
for s,n in [("3.1","เครื่องมือที่ใช้ในการพัฒนาระบบ"),("3.2","ขั้นตอนการดำเนินงาน"),
            ("3.3","การออกแบบการทำงานของระบบ"),("3.4","การกำหนดตัวแปรและโครงสร้างข้อมูล"),
            ("3.5","การออกแบบส่วนติดต่อประสานงานกับผู้ใช้"),("3.6","การออกแบบฐานข้อมูล")]:
    entry(newdoc, f"{s} {n}", 2, P_exact(s))
blank(newdoc)

entry(newdoc, "บทที่ 4 ผลการดำเนินงาน", 1, P_exact("บทที่ 4"))
for s,n in [("4.1","ผลของการพัฒนาระบบ"),("4.2","ผลของการทดสอบความถูกต้องของลิงก์และการนำทาง"),
            ("4.3","ผลการทดสอบการยอมรับของผู้ใช้งาน (UAT)"),("4.4","ผลการประเมินประสิทธิภาพและความพึงพอใจ")]:
    entry(newdoc, f"{s} {n}", 2, P_exact(s))
blank(newdoc)

entry(newdoc, "บทที่ 5 สรุปผล อภิปรายผล และข้อเสนอแนะ", 1, P_exact("บทที่ 5"))
for s,n in [("5.1","สรุปผลการดำเนินงาน"),("5.2","อภิปรายผล"),
            ("5.3","ปัญหาและอุปสรรค"),("5.4","ข้อเสนอแนะเพื่อการพัฒนาต่อยอด")]:
    entry(newdoc, f"{s} {n}", 2, P_exact(s))
blank(newdoc)
entry(newdoc, "บรรณานุกรม", 0, P_exact("บรรณานุกรม"))

# --- สารบัญตาราง ---
newdoc.add_page_break()
title(newdoc, "สารบัญตาราง")
hdr(newdoc, "ตาราง", "หน้า")
for t in tables:
    entry(newdoc, t, 2, P_exact(t))

# --- สารบัญภาพ ---
newdoc.add_page_break()
title(newdoc, "สารบัญภาพ")
hdr(newdoc, "ภาพ", "หน้า")
for f in figures:
    entry(newdoc, f, 2, P_exact(f))

out = r"C:\Users\tanut\Downloads\สารบัญ_NitiSmart_V4_Final.docx"
newdoc.save(out)
print(f"Saved: {out}")
