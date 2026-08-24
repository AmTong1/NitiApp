# -*- coding: utf-8 -*-
import glob
import os
from docx import Document
from docx.oxml.ns import qn

# Find the V1.9 file, exclude temp files (~$)
files = glob.glob(r"C:\Users\tanut\Downloads\*V1.9.docx")
files = [f for f in files if '~$' not in f]
if not files:
    print("ERROR: File not found!")
    exit(1)

filepath = files[0]
print(f"Reading: {filepath}")
print(f"File size: {os.path.getsize(filepath)} bytes")
print("="*80)

doc = Document(filepath)

# ---- Extract all headings, paragraphs, images, tables ----
print("\n### DOCUMENT STRUCTURE ###\n")

para_count = 0
table_count = 0
image_count = 0
heading_list = []
table_list = []
image_list = []

for i, element in enumerate(doc.element.body):
    tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
    
    if tag == 'p':
        para = None
        for p in doc.paragraphs:
            if p._element is element:
                para = p
                break
        if para:
            para_count += 1
            style_name = para.style.name if para.style else "None"
            text = para.text.strip()
            
            # Check for images in paragraph
            for run in para.runs:
                for drawing in run._element.findall(qn('w:drawing')):
                    image_count += 1
                    desc = ""
                    for docPr in drawing.iter(qn('wp:docPr')):
                        desc = docPr.get('descr', docPr.get('name', ''))
                    image_list.append((image_count, desc, style_name))
                    
            for pic in para._element.findall('.//' + qn('w:pict')):
                image_count += 1
                image_list.append((image_count, '', style_name))
            
            if 'Heading' in style_name or 'heading' in style_name:
                heading_list.append((style_name, text, para_count))
                print(f"[HEADING] Style={style_name} | Text=\"{text}\"")
            elif text and len(text) > 0:
                preview = text[:150] + ('...' if len(text) > 150 else '')
                if any(kw in text for kw in ['บทที่', 'สารบัญ', 'รูปที่', 'ตารางที่', 'ภาคผนวก', 'บรรณานุกรม', 'กิตติกรรม', 'บทคัดย่อ', 'Abstract', 'ABSTRACT', 'หน้า']):
                    print(f"  [KEY] Style={style_name} | Text=\"{preview}\"")
                    heading_list.append((style_name, text, para_count))
    
    elif tag == 'tbl':
        table_count += 1
        table_list.append((table_count, f"Table #{table_count}"))
        print(f"[TABLE] Table #{table_count}")

print("\n" + "="*80)
print(f"\nTotal paragraphs: {para_count}")
print(f"Total tables: {table_count}")
print(f"Total images: {image_count}")

print("\n\n### ALL HEADINGS ###\n")
for style, text, pnum in heading_list:
    print(f"  [{style}] {text}")

print("\n\n### FIGURE & TABLE CAPTIONS ###\n")
for para in doc.paragraphs:
    text = para.text.strip()
    style = para.style.name if para.style else "None"
    if text:
        if any(kw in text for kw in ['รูปที่', 'ตารางที่']):
            print(f"  [{style}] {text}")

print("\n\n### SECTIONS / PAGES ###\n")
# Check for section breaks and page numbering
for i, section in enumerate(doc.sections):
    print(f"Section {i}: width={section.page_width}, height={section.page_height}")
    if section.header:
        for p in section.header.paragraphs:
            if p.text.strip():
                print(f"  Header: {p.text.strip()}")
    if section.footer:
        for p in section.footer.paragraphs:
            if p.text.strip():
                print(f"  Footer: {p.text.strip()}")
