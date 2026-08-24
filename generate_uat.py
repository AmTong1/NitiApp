import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Cordia New'
        if level == 1:
            run.font.size = Pt(24)
        elif level == 2:
            run.font.size = Pt(20)

def create_uat_table(doc, items):
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ['Test ID', 'หมวดหมู่', 'ขั้นตอนการทดสอบ (Test Steps)', 'ผลลัพธ์ที่คาดหวัง (Expected Result)', 'ผล (Pass/Fail)']
    
    # Set headers
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Cordia New'
                run.font.size = Pt(16)
    
    # Add rows
    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item['id']
        row_cells[1].text = item['module']
        row_cells[2].text = item['steps']
        row_cells[3].text = item['expected']
        row_cells[4].text = "___ Pass\n___ Fail"
        
        for i in range(5):
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Cordia New'
                    run.font.size = Pt(14)

doc = docx.Document()
add_heading(doc, 'เอกสารทดสอบการยอมรับระบบ (UAT Test Script)', level=1)
doc.add_paragraph('ระบบบริหารจัดการนิติบุคคล (NitiSmart)').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

add_heading(doc, 'ส่วนที่ 1: การทดสอบฝั่งลูกบ้าน (Resident)', level=2)

resident_uat = [
    {
        'id': 'RES-01', 'module': 'Login & Profile',
        'steps': 'เข้าสู่ระบบด้วยเบอร์โทรศัพท์และรหัสผ่านที่ถูกต้อง',
        'expected': 'ระบบนำทางเข้าสู่หน้า Home (Dashboard) ได้สำเร็จ'
    },
    {
        'id': 'RES-02', 'module': 'Login & Profile',
        'steps': 'เข้าสู่หน้าตั้งค่า (Profile) แล้วทำการ "เปลี่ยนรหัสผ่าน"',
        'expected': 'เปลี่ยนรหัสสำเร็จ และสามารถใช้รหัสผ่านใหม่เข้าสู่ระบบได้'
    },
    {
        'id': 'RES-03', 'module': 'Dashboard',
        'steps': 'ตรวจสอบการแสดงผลหน้าแรก (Home)',
        'expected': 'แสดงยอดค้างชำระ (ถ้ามี), วันที่ครบกำหนด และประกาศล่าสุดได้อย่างถูกต้อง'
    },
    {
        'id': 'RES-04', 'module': 'Payment',
        'steps': 'ตรวจสอบเมนูประวัติการชำระเงิน',
        'expected': 'แสดงงวดที่ต้องชำระ (Pending), จำนวนงวดที่ชำระแล้ว และยอดรวมทั้งหมดได้ตรงตามจริง'
    },
    {
        'id': 'RES-05', 'module': 'Payment',
        'steps': 'กดปุ่มชำระเงินผ่าน QR Code',
        'expected': 'ระบบสร้างและแสดง QR Code สำหรับ PromptPay พร้อมระบุจำนวนเงินที่ต้องชำระได้อย่างถูกต้อง'
    },
    {
        'id': 'RES-06', 'module': 'Payment',
        'steps': 'อัปโหลดสลิปการโอนเงินและบันทึก',
        'expected': 'รายการชำระเงินเปลี่ยนสถานะเป็น "รออนุมัติ" และแสดงปุ่มดูภาพสลิปที่แนบไว้'
    },
    {
        'id': 'RES-07', 'module': 'Payment',
        'steps': 'กดดาวน์โหลดใบเสร็จรับเงินสำหรับงวดที่ชำระสำเร็จแล้ว',
        'expected': 'สามารถดาวน์โหลดและดูใบเสร็จรับเงินได้ โดยเวลาบนใบเสร็จตรงตามเวลาจริง'
    },
    {
        'id': 'RES-08', 'module': 'Repair',
        'steps': 'สร้างรายการแจ้งซ่อมใหม่ พร้อมระบุรายละเอียดและแนบรูปภาพ',
        'expected': 'รายการแจ้งซ่อมถูกสร้างและแสดงในหน้ารายการสถานะ "รอรับเรื่อง"'
    },
    {
        'id': 'RES-09', 'module': 'Repair',
        'steps': 'กดยกเลิกการแจ้งซ่อมของตนเองในรายการที่ยังไม่ดำเนินการ',
        'expected': 'สถานะการแจ้งซ่อมเปลี่ยนเป็น "ยกเลิก" สำเร็จ'
    },
    {
        'id': 'RES-10', 'module': 'Chat',
        'steps': 'ส่งข้อความติดต่อสอบถามไปยังนิติบุคคล',
        'expected': 'ข้อความปรากฏในห้องแชท และฝั่งนิติบุคคลได้รับข้อความ'
    },
    {
        'id': 'RES-11', 'module': 'Chat',
        'steps': 'ส่งไฟล์รูปภาพ, ไฟล์เอกสาร (PDF) หรือไฟล์ Excel ผ่านช่องแชท',
        'expected': 'ระบบสามารถแสดงผลรูปภาพ และสามารถคลิกเพื่อดาวน์โหลดไฟล์เอกสารได้ตามปกติ'
    },
    {
        'id': 'RES-12', 'module': 'Announcements',
        'steps': 'กดดูรายการประกาศจากนิติบุคคลในหน้าประกาศ',
        'expected': 'แสดงรายละเอียดของประกาศ รูปภาพประกอบ(ถ้ามี) ได้ครบถ้วน'
    }
]

create_uat_table(doc, resident_uat)
doc.add_page_break()

add_heading(doc, 'ส่วนที่ 2: การทดสอบฝั่งนิติบุคคล (Juristic Person / Admin)', level=2)

admin_uat = [
    {
        'id': 'ADM-01', 'module': 'Login',
        'steps': 'เข้าสู่ระบบด้วยสิทธิ์ Admin / SuperAdmin',
        'expected': 'สามารถเข้าสู่ระบบและเห็นเมนูการจัดการทั้งหมด (ลูกบ้าน, แจ้งซ่อม, บัญชี, แชท)'
    },
    {
        'id': 'ADM-02', 'module': 'Resident',
        'steps': 'เพิ่มข้อมูลบ้านและลูกบ้านใหม่ พร้อมระบุพื้นที่ (ตร.ม.)',
        'expected': 'ระบบบันทึกข้อมูลสำเร็จ สร้างรหัสผ่านเริ่มต้นให้ลูกบ้าน และสร้างงวดการชำระเงินอัตโนมัติ'
    },
    {
        'id': 'ADM-03', 'module': 'Resident',
        'steps': 'แก้ไขข้อมูลบ้าน และปรับเปลี่ยนขนาดพื้นที่ (ตร.ม.)',
        'expected': 'ข้อมูลอัปเดตสำเร็จ และระบบคำนวณค่าส่วนกลางใหม่ตามเรทพื้นที่ได้อย่างถูกต้อง'
    },
    {
        'id': 'ADM-04', 'module': 'Payment',
        'steps': 'ตรวจสอบและอนุมัติสลิปการชำระเงินของลูกบ้าน',
        'expected': 'สถานะงวดเปลี่ยนเป็น "ชำระแล้ว", บันทึกเวลาอนุมัติตรงตามเวลาจริง และสร้างใบเสร็จให้ลูกบ้าน'
    },
    {
        'id': 'ADM-05', 'module': 'Payment',
        'steps': 'ดูประวัติการทำรายการชำระเงิน (Logs)',
        'expected': 'แสดงประวัติการจ่ายเงิน การตรวจสอบสลิป และชื่อผู้อนุมัติ พร้อมเวลาที่ดำเนินการอย่างถูกต้อง'
    },
    {
        'id': 'ADM-06', 'module': 'Repair',
        'steps': 'รับเรื่องแจ้งซ่อมจากลูกบ้านในระบบ',
        'expected': 'สถานะของรายการแจ้งซ่อมเปลี่ยนเป็น "กำลังดำเนินการ"'
    },
    {
        'id': 'ADM-07', 'module': 'Repair',
        'steps': 'แก้ไขสถานะงานซ่อมเป็นปิดงาน พร้อมแนบรูปและรายละเอียดหลังซ่อมเสร็จ',
        'expected': 'สถานะการแจ้งซ่อมเปลี่ยนเป็น "เสร็จสิ้น" และฝั่งลูกบ้านมองเห็นการอัปเดตนี้'
    },
    {
        'id': 'ADM-08', 'module': 'Repair',
        'steps': 'ดูประวัติการดำเนินการงานซ่อม (Logs)',
        'expected': 'แสดงประวัติการรับเรื่อง, ปิดงาน รวมถึงการแก้ไขรายละเอียดโดยระบุตัวตน Admin ที่ดำเนินการ'
    },
    {
        'id': 'ADM-09', 'module': 'Financial',
        'steps': 'บันทึกรายการบัญชี (รายรับ-รายจ่าย) กำหนดเอง',
        'expected': 'ยอดรายรับ/รายจ่าย และยอดคงเหลือในระบบอัปเดตเรียลไทม์ และบันทึกประวัติสำเร็จ'
    },
    {
        'id': 'ADM-10', 'module': 'Financial',
        'steps': 'กดปุ่มเปิด/ปิด การมองเห็นยอดเงิน (Visibility) ให้ลูกบ้าน',
        'expected': 'สลับสถานะการเปิดให้ดูยอดเงินได้สำเร็จ (ลูกบ้านจะเห็นหรือไม่เห็นยอดรวมตามที่ Admin เปิด/ปิด)'
    },
    {
        'id': 'ADM-11', 'module': 'Financial',
        'steps': 'ส่งออกรายงานบัญชี (Export Excel)',
        'expected': 'สามารถดาวน์โหลดไฟล์ .xls รายการบัญชีและเปิดดูได้ พร้อมแสดงรายการรายรับ-รายจ่ายอย่างถูกต้อง'
    },
    {
        'id': 'ADM-12', 'module': 'Chat',
        'steps': 'ตอบแชทลูกบ้าน พร้อมส่งไฟล์รูปและไฟล์ Excel',
        'expected': 'สามารถส่งข้อความ/ไฟล์ได้ปกติ ลูกบ้านได้รับข้อความทันที (Real-time)'
    },
    {
        'id': 'ADM-13', 'module': 'Announcements',
        'steps': 'สร้างประกาศใหม่แบบมีความสำคัญ (Mark as Important)',
        'expected': 'ประกาศถูกสร้างขึ้น และแสดงผลบนหน้า Home ของลูกบ้านด้วยป้ายกำกับชัดเจน'
    },
    {
        'id': 'ADM-14', 'module': 'Announcements',
        'steps': 'ลบประกาศเก่าออกจากระบบ',
        'expected': 'ประกาศนั้นหายไปจากระบบทันที และลูกบ้านไม่สามารถเข้าดูได้อีก'
    }
]

create_uat_table(doc, admin_uat)

doc.save('NitiSmart_UAT_Test_Script.docx')
print('NitiSmart_UAT_Test_Script.docx generated successfully!')
